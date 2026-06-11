"""Контентные воронки: извлечение текста из файлов конкурентов, генерация
раскадровки слайдов через Claude API и экспорт брифа для дизайнера в HTML."""

import io
import json
import logging
import os
import re
from datetime import datetime
from html import escape

from anthropic import Anthropic

from wb_report import fmt_int, fmt_money, fmt_pct, short_money

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
ANTHROPIC_MODEL = os.environ.get("ANTHROPIC_FUNNEL_MODEL", "claude-sonnet-4-5")


def _client() -> Anthropic:
    if not ANTHROPIC_API_KEY:
        raise RuntimeError("ANTHROPIC_API_KEY не задан — добавьте переменную окружения в Railway")
    return Anthropic(api_key=ANTHROPIC_API_KEY)


SYSTEM_PROMPT = """Ты — маркетолог-аналитик, который составляет техническое задание (бриф)
для дизайнера на инфографику карточки товара Wildberries (контентную воронку).

На вход тебе дают JSON со следующими полями:
- own_data: собственные данные продавца — воронка продаж (показы, корзина, заказы, выкупы,
  CTR, конверсии), цена, юнит-экономика, остатки по складам;
- competitor_analysis: текстовая выжимка из отчётов по конкурентам (если есть);
- past_ab_learnings: выводы из прошлых А/Б тестов контента по этому товару (если есть);
- additional_notes: дополнительные пожелания пользователя.

Твоя задача — спроектировать последовательность слайдов (инфографик) для карточки товара:
обложка, ключевые УТП, закрытие болей и возражений покупателя, демонстрация использования,
сравнение с конкурентами, социальное доказательство, призыв к действию и т.д.

Для каждого слайда укажи:
- position — порядковый номер (1, 2, 3, ...);
- slide_type — тип слайда (например: cover, pain, solution, feature, social_proof,
  comparison, usage, cta — или другое значение, если ни одно не подходит);
- pain_point — какую боль/возражение покупателя закрывает этот слайд (если применимо,
  иначе пустая строка);
- message — главный посыл/смысл слайда (заголовок, который увидит покупатель);
- creative_notes — конкретные указания дизайнеру: что изобразить, какие акценты сделать,
  что написать мелким текстом, какой стиль/цвета использовать.

Учитывай слабые места из own_data (низкий CTR, низкая конверсия в корзину/заказ, высокая
цена и т.п.) и сильные стороны конкурентов из competitor_analysis. Если есть
past_ab_learnings — не повторяй неудачные решения и развивай удачные.

Ответь СТРОГО JSON-массивом объектов без пояснений и без markdown-обёртки, в формате:
[{"position": 1, "slide_type": "cover", "pain_point": "...", "message": "...", "creative_notes": "..."}, ...]
Сделай от 6 до 10 слайдов."""


def extract_text_from_upload(name: str, mime_type: str, file_data: bytes) -> str:
    """Извлекает текст из загруженного файла конкурента для передачи в промпт.

    HTML -> убираем <style>/<script>, get_text();
    PDF -> pypdf;
    text/* -> декодируем как есть;
    остальное (изображения и т.п.) -> "" (пока не передаются в промпт).
    Обрезка до 15000 символов.
    """
    text = ""
    lower_name = (name or "").lower()
    try:
        if mime_type == "application/pdf" or lower_name.endswith(".pdf"):
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_data))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif "html" in (mime_type or "") or lower_name.endswith((".html", ".htm")):
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(file_data.decode("utf-8", errors="ignore"), "html.parser")
            for tag in soup(["style", "script"]):
                tag.decompose()
            text = soup.get_text(separator="\n")
        elif (mime_type or "").startswith("text/") or lower_name.endswith((".txt", ".csv")):
            text = file_data.decode("utf-8", errors="ignore")
        else:
            return ""
    except Exception as e:
        logging.error(f"extract_text_from_upload {name}: {e}")
        return ""

    text = re.sub(r"\n\s*\n+", "\n\n", text).strip()
    return text[:15000]


def _parse_slides_json(text: str) -> list[dict]:
    """Устойчивый парсинг JSON — Claude может обернуть ответ в ```json fences."""
    m = re.search(r"\[.*\]", text, re.S)
    data = json.loads(m.group(0) if m else text)
    return [{
        "position": int(s.get("position", i + 1)),
        "slide_type": str(s.get("slide_type", ""))[:50],
        "pain_point": str(s.get("pain_point", "")),
        "message": str(s.get("message", "")),
        "creative_notes": str(s.get("creative_notes", "")),
    } for i, s in enumerate(data)]


def generate_funnel(own_data: dict, competitor_text: str, ab_context: str, notes: str = "") -> tuple[list[dict], dict]:
    """Вызывает Claude API и возвращает (список слайдов, сырой ответ для отладки)."""
    user_content = json.dumps({
        "own_data": own_data,
        "competitor_analysis": competitor_text[:20000],
        "past_ab_learnings": ab_context[:5000],
        "additional_notes": notes,
    }, ensure_ascii=False)

    client = _client()
    resp = client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=4096,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_content}],
    )
    text = "".join(b.text for b in resp.content if b.type == "text")
    slides = _parse_slides_json(text)
    raw = {
        "model": ANTHROPIC_MODEL,
        "raw_text": text,
        "usage": resp.usage.model_dump() if resp.usage else None,
    }
    return slides, raw


# ──────────────────────────────────────────────────────────────────────────────
# Экспорт брифа в HTML
# ──────────────────────────────────────────────────────────────────────────────

BRIEF_CSS = r"""
:root{--blue:#378ADD;--blue-dark:#0C447C;--blue-light:#E6F1FB;--blue-mid:#185FA5;--red:#E24B4A;--red-light:#FCEBEB;--red-dark:#A32D2D;--amber:#EF9F27;--amber-light:#FAEEDA;--amber-dark:#633806;--green:#639922;--green-light:#EAF3DE;--green-dark:#3B6D11;--gray-light:#F4F3F0;--border:#D3D1C7;--text:#1a1a1a;--text-sec:#5F5E5A}
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
body{font-family:'Golos Text',sans-serif;font-size:14px;color:var(--text);background:#F7F6F3;padding:24px 16px 48px}
.page{max-width:920px;margin:0 auto;background:#fff;border-radius:16px;padding:32px 36px 40px;border:.5px solid var(--border)}
h1{font-size:22px;font-weight:700;line-height:1.2;margin-bottom:4px}
.subtitle{font-size:13px;color:var(--text-sec);margin-bottom:20px}
hr{border:none;border-top:.5px solid var(--border);margin:20px 0}
.section-label{font-size:10px;font-weight:700;letter-spacing:.07em;text-transform:uppercase;color:#888780;margin:20px 0 10px}

.metrics-grid{display:grid;grid-template-columns:repeat(4,1fr);gap:10px;margin-bottom:14px}
.mc{background:var(--gray-light);border-radius:10px;padding:12px 14px}
.mc.hl{background:var(--red-light);border:.5px solid #F5C6C5}
.mc-label{font-size:10px;font-weight:700;color:var(--text-sec);letter-spacing:.05em;text-transform:uppercase;margin-bottom:4px}
.mc.hl .mc-label{color:var(--red-dark)}
.mc-val{font-size:22px;font-weight:700;line-height:1.1;margin-bottom:3px}
.mc.hl .mc-val{color:var(--red-dark)}
.mc-delta{font-size:12px}
.du{color:var(--green-dark)}.dd{color:var(--red-dark)}.dn{color:var(--text-sec)}

.legend{display:flex;flex-direction:column;gap:4px;margin-bottom:14px}
.legend-row{display:flex;align-items:center;gap:8px;font-size:12px}
.legend-badge{font-size:10px;font-weight:700;padding:2px 8px;border-radius:4px;white-space:nowrap}
.legend-desc{color:var(--text-sec)}
.tasks{display:flex;flex-direction:column;border:.5px solid var(--border);border-radius:12px;overflow:hidden}
.task{padding:10px 14px 12px;border-bottom:.5px solid var(--border);border-left:3px solid transparent;transition:background .15s}
.task:last-child{border-bottom:none}
.task:hover{background:#FAFAF8}
.task-crit{border-left-color:var(--red)}.task-mid{border-left-color:var(--amber)}.task-good{border-left-color:var(--green)}
.task-header{display:flex;align-items:center;gap:8px;margin-bottom:5px}
.task-badge{font-size:10px;font-weight:700;padding:2px 8px;border-radius:4px;white-space:nowrap}
.badge-crit{background:var(--red-light);color:var(--red-dark)}
.badge-mid{background:var(--amber-light);color:var(--amber-dark)}
.badge-good{background:var(--green-light);color:var(--green-dark)}
.badge-info{background:var(--blue-light);color:var(--blue-dark)}
.task-tag{font-size:10px;font-weight:700;color:var(--text-sec);letter-spacing:.05em;text-transform:uppercase}
.task-title{font-size:13px;font-weight:600;margin-bottom:3px}
.task-body{font-size:12px;color:var(--text-sec);line-height:1.6;white-space:pre-line}

.cmp-wrap{overflow-x:auto;margin-top:4px}
table.cmp{width:100%;border-collapse:collapse;font-size:12px}
table.cmp th{font-size:10px;font-weight:700;color:var(--text-sec);letter-spacing:.04em;text-transform:uppercase;padding:6px 8px;border-bottom:.5px solid var(--border);text-align:left;background:var(--gray-light);white-space:nowrap}
table.cmp td{padding:7px 8px;border-bottom:.5px solid var(--border);color:var(--text);vertical-align:middle;white-space:nowrap}
table.cmp tr:last-child td{border-bottom:none}
table.cmp tr:hover td{background:#FAFAF8}
.pill{display:inline-flex;align-items:center;font-size:10px;font-weight:700;padding:2px 7px;border-radius:20px;white-space:nowrap}
.pill-win{background:var(--green-light);color:var(--green-dark)}
.pill-lose{background:var(--red-light);color:var(--red-dark)}
.pill-tie{background:var(--blue-light);color:var(--blue-dark)}

.wh-summary{display:grid;grid-template-columns:repeat(4,1fr);gap:10px;margin-bottom:14px}
.wh-sc{background:var(--gray-light);border-radius:10px;padding:12px;text-align:center}
.wh-sc .v{font-size:20px;font-weight:700;color:var(--text)}
.wh-sc .l{font-size:10px;color:var(--text-sec);text-transform:uppercase;letter-spacing:.04em;margin-top:3px}
table.wh{width:100%;border-collapse:collapse;font-size:12px}
table.wh th{font-size:10px;font-weight:700;color:var(--text-sec);letter-spacing:.04em;text-transform:uppercase;padding:6px 8px;border-bottom:.5px solid var(--border);text-align:left;background:var(--gray-light)}
table.wh td{padding:7px 8px;border-bottom:.5px solid var(--border);color:var(--text)}
table.wh tr:last-child td{border-bottom:none}
table.wh td.num{text-align:right;font-variant-numeric:tabular-nums;font-family:'JetBrains Mono',monospace}
table.wh tr.total td{font-weight:700;background:var(--gray-light)}
.note{font-size:11px;color:var(--text-sec);margin-top:8px;font-style:italic;line-height:1.5;white-space:pre-line}

@media(max-width:640px){.metrics-grid,.wh-summary{grid-template-columns:repeat(2,1fr)}}
"""


def _mc_delta(cur, prev) -> str:
    if not prev:
        return '<div class="mc-delta dn">нет данных за пред. период</div>'
    pct = (cur - prev) / prev * 100
    arrow = "▲" if pct > 0 else ("▼" if pct < 0 else "■")
    cls = "du" if pct > 0 else ("dd" if pct < 0 else "dn")
    return f'<div class="mc-delta {cls}">{arrow} {abs(pct):.0f}% к пред. периоду</div>'


def _delta_text(cur, prev) -> str:
    if not prev:
        return "нет данных за пред. период"
    pct = (cur - prev) / prev * 100
    arrow = "▲" if pct > 0 else ("▼" if pct < 0 else "■")
    return f"{arrow} {abs(pct):.0f}% к пред. периоду"


def _mc(label: str, val: str, delta_html: str = "") -> str:
    return (
        f'<div class="mc"><div class="mc-label">{escape(label)}</div>'
        f'<div class="mc-val">{escape(val)}</div>{delta_html}</div>'
    )


SLIDE_TYPE_LABELS = {
    "cover": "Обложка",
    "pain": "Боль",
    "solution": "Решение",
    "feature": "УТП / Особенность",
    "social_proof": "Соц. доказательство",
    "comparison": "Сравнение",
    "usage": "Применение",
    "cta": "Призыв к действию",
}


def render_brief_html(funnel: dict, slides: list[dict], own_data: dict) -> str:
    """Самодостаточный HTML-отчёт (бриф) для дизайнера в стиле wb_analysis_*.html."""
    project = funnel.get("project", "")
    wb_article = funnel.get("wb_article", "")
    seller_article = funnel.get("seller_article") or ""
    title = funnel.get("title") or f"Артикул {wb_article}"
    generated_at = datetime.now().strftime("%d.%m.%Y %H:%M")

    parts = []
    parts.append(f"<h1>Бриф на контентную воронку · {escape(title)}</h1>")
    subtitle = f"Кабинет: {escape(project)} · Артикул WB: {escape(str(wb_article))}"
    if seller_article:
        subtitle += f" · Артикул продавца: {escape(seller_article)}"
    subtitle += f" · Сформировано: {generated_at}"
    parts.append(f'<p class="subtitle">{subtitle}</p>')
    parts.append("<hr>")

    funnel_data = own_data.get("funnel") or {}
    price_data = own_data.get("price") or {}
    stocks_data = own_data.get("stocks") or {}
    unit_data = own_data.get("unit_economics") or {}

    if funnel_data or price_data:
        parts.append('<div class="section-label">Свои данные · воронка продаж (30 дней)</div>')
        cards = []
        if funnel_data:
            cards.append(_mc("Показы", fmt_int(funnel_data.get("views")),
                              _mc_delta(funnel_data.get("views"), funnel_data.get("views_prev"))))
            cards.append(_mc("CTR (корзина/показы)", fmt_pct(funnel_data.get("ctr_pct"))))
            cards.append(_mc("Заказы", fmt_int(funnel_data.get("orders")),
                              _mc_delta(funnel_data.get("orders"), funnel_data.get("orders_prev"))))
            cards.append(_mc("% выкупа", fmt_pct(funnel_data.get("buyout_pct"))))
            cards.append(_mc("Корзина", fmt_int(funnel_data.get("cart")),
                              _mc_delta(funnel_data.get("cart"), funnel_data.get("cart_prev"))))
            cards.append(_mc("Конв. в заказ из корзины", fmt_pct(funnel_data.get("cart_to_order_pct"))))
            cards.append(_mc("Оборот заказов", fmt_money(funnel_data.get("order_sum"))))
            cards.append(_mc("Выкупы", fmt_int(funnel_data.get("buyouts")),
                              _mc_delta(funnel_data.get("buyouts"), funnel_data.get("buyouts_prev"))))
        if price_data.get("discounted_price") is not None:
            cards.append(_mc("Цена со скидкой", fmt_money(price_data.get("discounted_price"))))
        if stocks_data.get("total") is not None:
            cards.append(_mc("Остатки на складах", f"{fmt_int(stocks_data.get('total'))} шт"))
        parts.append(f'<div class="metrics-grid">{"".join(cards)}</div>')

    if unit_data:
        parts.append('<div class="section-label">Юнит-экономика</div>')
        cards = []
        if unit_data.get("cost_price") is not None:
            cards.append(_mc("Себестоимость", fmt_money(unit_data.get("cost_price"))))
        if unit_data.get("commission_pct") is not None:
            cards.append(_mc("Комиссия WB", fmt_pct(unit_data.get("commission_pct"))))
        if unit_data.get("redemption_pct") is not None:
            cards.append(_mc("% выкупа (юнит-эконом.)", fmt_pct(unit_data.get("redemption_pct"))))
        if unit_data.get("stock") is not None:
            cards.append(_mc("Остаток (юнит-эконом.)", f"{fmt_int(unit_data.get('stock'))} шт"))
        if cards:
            parts.append(f'<div class="metrics-grid">{"".join(cards)}</div>')

    by_warehouse = stocks_data.get("by_warehouse") or []
    if by_warehouse:
        parts.append("<hr>")
        parts.append('<div class="section-label">Остатки по складам</div>')
        rows = "".join(
            f'<tr><td>{escape(w["warehouse"])}</td><td class="num">{fmt_int(w["qty"])}</td></tr>'
            for w in by_warehouse
        )
        total = fmt_int(stocks_data.get("total"))
        parts.append(
            '<div class="cmp-wrap"><table class="wh"><thead><tr>'
            '<th>Склад</th><th style="text-align:right">Остаток, шт</th>'
            f'</tr></thead><tbody>{rows}'
            f'<tr class="total"><td>Итого</td><td class="num">{total}</td></tr>'
            '</tbody></table></div>'
        )

    parts.append("<hr>")
    parts.append('<div class="section-label">Контентная воронка · слайды карточки</div>')
    if slides:
        parts.append('<div class="tasks">')
        for s in slides:
            slide_type = s.get("slide_type", "")
            badge_label = SLIDE_TYPE_LABELS.get(slide_type, slide_type or "Слайд")
            position = s.get("position", "")
            title_text = s.get("message") or s.get("pain_point") or ""
            body_lines = []
            if s.get("pain_point") and s.get("message"):
                body_lines.append(f"Боль/смысл: {s['pain_point']}")
            if s.get("creative_notes"):
                body_lines.append(f"Креатив: {s['creative_notes']}")
            body = "\n".join(body_lines)
            parts.append(
                '<div class="task">'
                '<div class="task-header">'
                f'<span class="task-badge badge-info">{escape(badge_label)}</span>'
                f'<span class="task-tag">Слайд {escape(str(position))}</span>'
                '</div>'
                f'<div class="task-title">{escape(title_text)}</div>'
                f'<div class="task-body">{escape(body)}</div>'
                '</div>'
            )
        parts.append("</div>")
    else:
        parts.append('<p class="note">Слайды ещё не заполнены.</p>')

    ab_context = (funnel.get("ab_context") or "").strip()
    if ab_context:
        parts.append("<hr>")
        parts.append('<div class="section-label">Учтены прошлые А/Б тесты</div>')
        parts.append(f'<p class="note">{escape(ab_context)}</p>')

    body_html = "\n".join(parts)
    return f"""<!DOCTYPE html>
<html lang="ru">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Бриф · {escape(title)}</title>
<link href="https://fonts.googleapis.com/css2?family=Golos+Text:wght@400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap" rel="stylesheet">
<style>{BRIEF_CSS}</style>
</head>
<body>
<div class="page">
{body_html}
</div>
</body>
</html>"""


# ──────────────────────────────────────────────────────────────────────────────
# Экспорт брифа в PDF / Word
# ──────────────────────────────────────────────────────────────────────────────

FONTS_DIR = os.path.join(os.path.dirname(__file__), "fonts")


def _brief_sections(funnel: dict, slides: list[dict], own_data: dict) -> dict:
    """Общие данные для PDF/Word-брифа: заголовок, метрики, слайды."""
    project = funnel.get("project", "")
    wb_article = funnel.get("wb_article", "")
    seller_article = funnel.get("seller_article") or ""
    title = funnel.get("title") or f"Артикул {wb_article}"

    subtitle = f"Кабинет: {project} · Артикул WB: {wb_article}"
    if seller_article:
        subtitle += f" · Артикул продавца: {seller_article}"
    subtitle += f" · Сформировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}"

    funnel_data = own_data.get("funnel") or {}
    price_data = own_data.get("price") or {}
    stocks_data = own_data.get("stocks") or {}
    unit_data = own_data.get("unit_economics") or {}

    funnel_rows = []
    if funnel_data:
        funnel_rows.append(("Показы", f"{fmt_int(funnel_data.get('views'))}  ({_delta_text(funnel_data.get('views'), funnel_data.get('views_prev'))})"))
        funnel_rows.append(("CTR (корзина/показы)", fmt_pct(funnel_data.get("ctr_pct"))))
        funnel_rows.append(("Корзина", f"{fmt_int(funnel_data.get('cart'))}  ({_delta_text(funnel_data.get('cart'), funnel_data.get('cart_prev'))})"))
        funnel_rows.append(("Конв. в заказ из корзины", fmt_pct(funnel_data.get("cart_to_order_pct"))))
        funnel_rows.append(("Заказы", f"{fmt_int(funnel_data.get('orders'))}  ({_delta_text(funnel_data.get('orders'), funnel_data.get('orders_prev'))})"))
        funnel_rows.append(("Оборот заказов", fmt_money(funnel_data.get("order_sum"))))
        funnel_rows.append(("Выкупы", f"{fmt_int(funnel_data.get('buyouts'))}  ({_delta_text(funnel_data.get('buyouts'), funnel_data.get('buyouts_prev'))})"))
        funnel_rows.append(("% выкупа", fmt_pct(funnel_data.get("buyout_pct"))))
    if price_data.get("discounted_price") is not None:
        funnel_rows.append(("Цена со скидкой", fmt_money(price_data.get("discounted_price"))))
    if stocks_data.get("total") is not None:
        funnel_rows.append(("Остатки на складах", f"{fmt_int(stocks_data.get('total'))} шт"))

    unit_rows = []
    if unit_data.get("cost_price") is not None:
        unit_rows.append(("Себестоимость", fmt_money(unit_data.get("cost_price"))))
    if unit_data.get("commission_pct") is not None:
        unit_rows.append(("Комиссия WB", fmt_pct(unit_data.get("commission_pct"))))
    if unit_data.get("redemption_pct") is not None:
        unit_rows.append(("% выкупа (юнит-эконом.)", fmt_pct(unit_data.get("redemption_pct"))))
    if unit_data.get("stock") is not None:
        unit_rows.append(("Остаток (юнит-эконом.)", f"{fmt_int(unit_data.get('stock'))} шт"))

    by_warehouse = stocks_data.get("by_warehouse") or []
    warehouse_rows = [(str(w.get("warehouse", "")), fmt_int(w.get("qty"))) for w in by_warehouse]
    warehouse_total = fmt_int(stocks_data.get("total")) if by_warehouse else None

    slide_items = []
    for s in slides:
        slide_type = s.get("slide_type", "")
        slide_items.append({
            "position": s.get("position", ""),
            "badge": SLIDE_TYPE_LABELS.get(slide_type, slide_type or "Слайд"),
            "title": s.get("message") or s.get("pain_point") or "",
            "pain_point": s.get("pain_point") or "",
            "message": s.get("message") or "",
            "creative_notes": s.get("creative_notes") or "",
        })

    return {
        "title": title,
        "subtitle": subtitle,
        "funnel_rows": funnel_rows,
        "unit_rows": unit_rows,
        "warehouse_rows": warehouse_rows,
        "warehouse_total": warehouse_total,
        "slides": slide_items,
        "ab_context": (funnel.get("ab_context") or "").strip(),
    }


def render_brief_pdf(funnel: dict, slides: list[dict], own_data: dict) -> bytes:
    """Бриф для дизайнера в виде PDF (reportlab + шрифт DejaVu Sans для кириллицы)."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    if "Body" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("Body", os.path.join(FONTS_DIR, "DejaVuSans.ttf")))
        pdfmetrics.registerFont(TTFont("Body-Bold", os.path.join(FONTS_DIR, "DejaVuSans-Bold.ttf")))

    GRAY = colors.HexColor("#5F5E5A")
    BORDER = colors.HexColor("#D3D1C7")
    BG = colors.HexColor("#F4F3F0")

    style_title = ParagraphStyle("title", fontName="Body-Bold", fontSize=16, leading=20)
    style_sub = ParagraphStyle("sub", fontName="Body", fontSize=9, leading=13, textColor=GRAY, spaceAfter=10)
    style_h2 = ParagraphStyle("h2", fontName="Body-Bold", fontSize=10, leading=14, textColor=GRAY,
                               spaceBefore=14, spaceAfter=6)
    style_body = ParagraphStyle("body", fontName="Body", fontSize=9, leading=13)
    style_bold = ParagraphStyle("bold", fontName="Body-Bold", fontSize=9, leading=13)
    style_note = ParagraphStyle("note", fontName="Body", fontSize=8.5, leading=12.5, textColor=GRAY)

    sec = _brief_sections(funnel, slides, own_data)
    elements = [
        Paragraph(escape(f"Бриф на контентную воронку · {sec['title']}"), style_title),
        Paragraph(escape(sec["subtitle"]), style_sub),
    ]

    def metric_table(rows, col_widths=(95 * mm, 75 * mm)):
        data = [[Paragraph(escape(label), style_body), Paragraph(escape(value), style_bold)] for label, value in rows]
        t = Table(data, colWidths=list(col_widths))
        t.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, BORDER),
            ("BACKGROUND", (0, 0), (0, -1), BG),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        return t

    if sec["funnel_rows"]:
        elements.append(Paragraph("Свои данные · воронка продаж (30 дней)", style_h2))
        elements.append(metric_table(sec["funnel_rows"]))

    if sec["unit_rows"]:
        elements.append(Paragraph("Юнит-экономика", style_h2))
        elements.append(metric_table(sec["unit_rows"]))

    if sec["warehouse_rows"]:
        elements.append(Paragraph("Остатки по складам", style_h2))
        data = [[Paragraph("Склад", style_bold), Paragraph("Остаток, шт", style_bold)]]
        for wh, qty in sec["warehouse_rows"]:
            data.append([Paragraph(escape(wh), style_body), Paragraph(escape(qty), style_body)])
        data.append([Paragraph("Итого", style_bold), Paragraph(escape(sec["warehouse_total"] or ""), style_bold)])
        t = Table(data, colWidths=[110 * mm, 60 * mm])
        t.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, BORDER),
            ("BACKGROUND", (0, 0), (-1, 0), BG),
            ("BACKGROUND", (0, -1), (-1, -1), BG),
            ("ALIGN", (1, 0), (1, -1), "RIGHT"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        elements.append(t)

    elements.append(Paragraph("Контентная воронка · слайды карточки", style_h2))
    if sec["slides"]:
        for s in sec["slides"]:
            elements.append(Paragraph(escape(f"Слайд {s['position']} · {s['badge']}"), style_bold))
            if s["title"]:
                elements.append(Paragraph(escape(s["title"]), style_body))
            if s["pain_point"] and s["message"]:
                elements.append(Paragraph(f"Боль/смысл: {escape(s['pain_point'])}", style_note))
            if s["creative_notes"]:
                note = escape(s["creative_notes"]).replace("\n", "<br/>")
                elements.append(Paragraph(f"Креатив: {note}", style_note))
            elements.append(Spacer(1, 8))
    else:
        elements.append(Paragraph("Слайды ещё не заполнены.", style_note))

    if sec["ab_context"]:
        elements.append(Paragraph("Учтены прошлые А/Б тесты", style_h2))
        elements.append(Paragraph(escape(sec["ab_context"]).replace("\n", "<br/>"), style_note))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm, topMargin=16 * mm, bottomMargin=16 * mm,
        title=f"Бриф · {sec['title']}",
    )
    doc.build(elements)
    return buf.getvalue()


def render_brief_docx(funnel: dict, slides: list[dict], own_data: dict) -> bytes:
    """Бриф для дизайнера в виде Word-документа (.docx)."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    GRAY = RGBColor(0x5F, 0x5E, 0x5A)
    sec = _brief_sections(funnel, slides, own_data)

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    doc.add_heading(f"Бриф на контентную воронку · {sec['title']}", level=1)
    p = doc.add_paragraph(sec["subtitle"])
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = GRAY

    def add_table(rows):
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value
        doc.add_paragraph()

    if sec["funnel_rows"]:
        doc.add_heading("Свои данные · воронка продаж (30 дней)", level=2)
        add_table(sec["funnel_rows"])

    if sec["unit_rows"]:
        doc.add_heading("Юнит-экономика", level=2)
        add_table(sec["unit_rows"])

    if sec["warehouse_rows"]:
        doc.add_heading("Остатки по складам", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Склад"
        hdr[1].text = "Остаток, шт"
        for wh, qty in sec["warehouse_rows"]:
            cells = table.add_row().cells
            cells[0].text = wh
            cells[1].text = qty
        cells = table.add_row().cells
        cells[0].text = "Итого"
        cells[1].text = sec["warehouse_total"] or ""
        doc.add_paragraph()

    doc.add_heading("Контентная воронка · слайды карточки", level=2)
    if sec["slides"]:
        for s in sec["slides"]:
            p = doc.add_paragraph()
            p.add_run(f"Слайд {s['position']} · {s['badge']}").bold = True
            if s["title"]:
                doc.add_paragraph(s["title"])
            if s["pain_point"] and s["message"]:
                p = doc.add_paragraph(f"Боль/смысл: {s['pain_point']}")
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = GRAY
            if s["creative_notes"]:
                p = doc.add_paragraph(f"Креатив: {s['creative_notes']}")
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = GRAY
    else:
        doc.add_paragraph("Слайды ещё не заполнены.")

    if sec["ab_context"]:
        doc.add_heading("Учтены прошлые А/Б тесты", level=2)
        doc.add_paragraph(sec["ab_context"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
